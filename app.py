import os
import faiss
import traceback
import traceback  # For debugging errors
import numpy as np
import streamlit as st
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEndpoint
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.docstore.in_memory import InMemoryDocstore
# from langchain.prompts import PromptTemplate

from secret_api_keys import huggingface_api_key  # Set the Hugging Face Hub API token as an environment variable


os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key

def process_input(input_type, input_data):
    """
    Processes different types of input data (Link, PDF, Text, DOCX, TXT) 
    and converts them into a FAISS-based vector store for efficient retrieval.

    Parameters:
    -----------
    input_type : str
        The type of input data. Supported types: "Link", "PDF", "Text", "DOCX", "TXT".
    input_data : str, BytesIO, or UploadedFile
        The actual content to be processed, which could be a URL (for "Link"), 
        a PDF file, a text string, a DOCX file, or a TXT file.

    Returns:
    --------
    FAISS
        A FAISS vector store containing the embedded representations of the processed input data.

    Raises:
    -------
    ValueError
        If an unsupported input type is provided or if the input data format is invalid.
    """
   
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data  # Input is already a text string
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [ str(doc.page_content) for doc in texts ]  # Access page_content from each Document 
    else:
        texts = text_splitter.split_text(documents)


    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )


    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)


    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )


    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store

def answer_question(vectorstore, query):
    """
    Generates an answer to a given query by retrieving relevant information 
    from a FAISS vector store and using a language model for response generation.

    Parameters:
    -----------
    vectorstore : FAISS
        The FAISS vector store containing embedded text representations for retrieval.
    query : str
        The question to be answered.

    Returns:
    --------
    dict
        A dictionary containing the generated answer.

    Raises:
    -------
    ValueError
        If the query is not a string.
    """

    
    llm = HuggingFaceEndpoint(repo_id= 'meta-llama/Meta-Llama-3-8B-Instruct', 
                              token = huggingface_api_key, temperature= 0.6)
    
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())

    answer = qa({"query": query})

    return answer



# Initialize session state for chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

def main():
    """
    Runs the RAG Chatbot application using Streamlit.

    The application allows users to select an input type (Link, PDF, Text, DOCX, TXT), 
    provide corresponding input data, and process it into a FAISS vector store. 
    Users can then ask questions based on the processed data, and the chatbot 
    retrieves relevant information to generate responses.

    Functionality:
    --------------
    - Allows users to select an input type and provide input data.
    - Supports URL inputs, text inputs, and file uploads (PDF, DOCX, TXT).
    - Processes the input data into a FAISS vector store.
    - Enables conversational interaction with the chatbot.
    - Displays chat history and handles errors gracefully.

    Raises:
    -------
    ValueError
        If the input type is invalid or no valid input data is provided.
    Exception
        If an error occurs during input processing or chatbot interaction.
    """


    st.title("💬 RAG Chatbot")

    # Placeholder option for input type selection
    input_type = st.selectbox(
        "Select Input Type",
        ["Select an Input Type", "Link", "PDF", "Text", "DOCX", "TXT"],
        index=0
    )

    input_data = None  # Default value

    try:
        if input_type == "Select an Input Type":
            st.warning("Please select an input type to proceed.")

        else:
            # Handling different input types
            if input_type == "Link":
                number_input = st.number_input("Enter the number of Links", min_value=1, max_value=20, step=1)
                input_data = []
                for i in range(number_input):
                    url = st.sidebar.text_input(f"URL {i+1}")
                    if url.strip():  # Ensure the URL is not empty
                        input_data.append(url)
                if not input_data:
                    st.warning("Please enter at least one valid URL.")

            elif input_type == "Text":
                input_data = st.text_area("Enter the text")  # Allows multiline input
                if not input_data.strip():
                    st.warning("Text input cannot be empty.")

            elif input_type in ["PDF", "TXT", "DOCX"]:
                input_data = st.file_uploader(f"Upload a {input_type} file", type=[input_type.lower()])
                if input_data is None:
                    st.warning(f"Please upload a valid {input_type} file.")

            # Proceed button
            if st.button("Process Data"):
                if not input_data:
                    st.error("No valid input provided. Please enter text, a URL, or upload a file.")
                else:
                    try:
                        vectorstore = process_input(input_type, input_data)
                        if vectorstore is None:
                            st.error("Failed to create the vector store. Please check the input data.")
                        else:
                            st.session_state["vectorstore"] = vectorstore
                            st.success("Vector store successfully created!")
                    except Exception as e:
                        st.error(f"Error processing input: {str(e)}")
                        st.text(traceback.format_exc())  # Show traceback for debugging

        # Display chat history
        st.subheader("Chatbot Conversation")
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        # Chat input field
        query = st.chat_input("Ask me a question...")  # Chatbot-style input

        if query:
            st.session_state.messages.append({"role": "user", "content": query})  # Store user question

            if "vectorstore" in st.session_state:
                try:
                    response = answer_question(st.session_state["vectorstore"], query)

                    # Extract only the answer text
                    if isinstance(response, dict) and "result" in response:
                        answer_text = response["result"]
                    else:
                        answer_text = "I'm not sure about that. Try another question!"

                    st.session_state.messages.append({"role": "assistant", "content": answer_text})  # Store bot response
                    st.rerun()  # Refresh chat UI

                except Exception as e:
                    st.error(f"Error retrieving answer: {str(e)}")
                    st.text(traceback.format_exc())

    except Exception as e:
        st.error(f"Unexpected error: {str(e)}")
        st.text(traceback.format_exc())  # Display error traceback for debugging

if __name__ == "__main__":
    main()

